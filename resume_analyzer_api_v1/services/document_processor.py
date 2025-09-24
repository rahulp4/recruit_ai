# document_processor.py

import pandas as pd
from docx import Document
import io
import logging

# --- Logging Setup ---
# Set up a logger for this module
logger = logging.getLogger(__name__)
# It's good practice for a library module not to configure basicConfig
# as it might interfere with the main app's logging setup.
# If you run this file directly for testing, you might temporarily uncomment:
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')


class DocumentProcessor:
    """
    A class to process DOCX documents, extracting both
    structured table data and general paragraph text.
    """
    def __init__(self, docx_content_stream: io.BytesIO):
        """
        Initializes the DocumentProcessor with a DOCX file stream.
        Loads the DOCX document once for all subsequent operations.

        Args:
            docx_content_stream (io.BytesIO): A stream containing the byte content of the DOCX file.
        """
        if not isinstance(docx_content_stream, io.BytesIO):
            raise TypeError("docx_content_stream must be an io.BytesIO object.")
        
        try:
            docx_content_stream.seek(0)  # Ensure stream is at the start
            self._document = Document(docx_content_stream)
            logger.info("DocumentProcessor initialized: DOCX document loaded successfully.")
        except Exception as e:
            logger.error(f"Error initializing DocumentProcessor: Could not load DOCX document: {e}", exc_info=True)
            raise # Re-raise to indicate initialization failure

    def _extract_tables_from_docx(self) -> list:
        """
        (Internal method) Extracts tables from the loaded DOCX document.

        Returns:
            list: A list of dictionaries, where each dictionary contains:
                  'table_index' (int) and 'dataframe' (pd.DataFrame) for an extracted table.
        """
        all_extracted_dfs = []
        if not self._document.tables:
            logger.info("No tables found in the document.")
            return []

        for table_idx, table in enumerate(self._document.tables):
            logger.info(f"Processing Table {table_idx + 1}...")
            data = []
            keys = None

            for i, row in enumerate(table.rows):
                text_cells = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    keys = text_cells
                else:
                    data.append(text_cells)

            if keys:
                if data:
                    max_cols = len(keys)
                    padded_data = [row + [None] * (max_cols - len(row)) if len(row) < max_cols else row[:max_cols] for row in data]
                    df = pd.DataFrame(padded_data, columns=keys)
                else:
                    df = pd.DataFrame(columns=keys)

                all_extracted_dfs.append({
                    "table_index": table_idx,
                    "dataframe": df
                })
                logger.info(f"  Extracted Table {table_idx + 1}.")
            else:
                logger.warning(f"  Table {table_idx + 1} has no header row or no content.")
        return all_extracted_dfs

    def _extract_text_from_docx(self) -> str:
        """
        (Internal method) Extracts all textual content from top-level paragraphs
        in the loaded DOCX document (excluding tables).

        Returns:
            str: The concatenated text from paragraphs.
        """
        full_text = []
        for para in self._document.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        return "\n".join(full_text)

    def _format_dataframe_rows_to_sections(self, df: pd.DataFrame, date_column_name: str = 'Date Range') -> list:
        """
        (Internal method) Formats each row of a DataFrame into a structured "block of sections."

        Args:
            df (pd.DataFrame): The DataFrame containing the tabular data.
            date_column_name (str): The name of the column containing the date range.

        Returns:
            list: A list of strings, where each string is a formatted section block.
        """
        formatted_sections = []
        if df.empty:
            logger.info("  DataFrame is empty, no sections to format.")
            return []

        if date_column_name not in df.columns:
            logger.warning(f"  Date column '{date_column_name}' not found. Looking for alternatives.")
            possible_date_cols = ['Date', 'Period', 'Duration', 'Time Frame', 'Valid From', 'Invoice Date', 'Start Date']
            found_date_col = None
            for col in possible_date_cols:
                if col in df.columns:
                    date_column_name = col
                    found_date_col = True
                    logger.info(f"  Using '{date_column_name}' as the date column.")
                    break
            if not found_date_col:
                logger.warning("  Could not find a suitable date column. 'Date Range' will be 'N/A Date Range'.")
                date_column_name = None

        for index, row in df.iterrows():
            if date_column_name and date_column_name in row:
                date_range = str(row[date_column_name]).strip() if pd.notna(row[date_column_name]) else "N/A Date Range"
            else:
                date_range = "N/A Date Range"

            other_details = []
            for col_name, value in row.items():
                if col_name != date_column_name and col_name:
                    clean_value = str(value).strip() if pd.notna(value) else "N/A"
                    other_details.append(f"**{col_name}:** {clean_value}")

            details_block = "\n".join(other_details)

            section_block = f"""
---
## Period: {date_range}

{details_block}

"""
            formatted_sections.append(section_block)
        return formatted_sections

    def get_combined_document_content(self, date_column_name: str = 'Date Range') -> str:
        """
        Orchestrates the extraction of both paragraph text and formatted table data,
        combining them into a single string. This is the primary public method.

        Args:
            date_column_name (str): The name of the column in tables that contains the date range.

        Returns:
            str: The combined string of paragraph text followed by formatted table sections.
        """
        logger.info("Starting get_combined_document_content for comprehensive content extraction.")

        full_document_paragraph_text = ""
        formatted_tables_text = ""
        combined_document_content = ""

        # --- Step 1: Extract General Paragraph Text ---
        try:
            full_document_paragraph_text = self._extract_text_from_docx()
            logger.info(f"Debug: Paragraph text length: {len(full_document_paragraph_text)}")
            if not full_document_paragraph_text.strip():
                logger.warning("Debug: Paragraph text content is empty or only whitespace.")
        except Exception as e:
            logger.error(f"Failed to extract full paragraph text from DOCX: {e}", exc_info=True)

        # --- Step 2: Extract and Process Tables ---
        extracted_tables_info = self._extract_tables_from_docx()

        if not extracted_tables_info:
            logger.info("No tables were successfully extracted.")
        else:
            logger.info(f"Successfully extracted {len(extracted_tables_info)} tables.")
            all_formatted_sections = []

            for table_info in extracted_tables_info:
                current_df = table_info['dataframe']
                logger.debug(f"Processing Table {table_info['table_index'] + 1} for Section Formatting.")
                if current_df.empty:
                    logger.debug(f"Table {table_info['table_index'] + 1} is empty, skipping section formatting.")
                    continue

                formatted_table_sections = self._format_dataframe_rows_to_sections(
                    current_df,
                    date_column_name=date_column_name
                )
                all_formatted_sections.extend(formatted_table_sections)
                logger.debug(f"Generated {len(formatted_table_sections)} sections for Table {table_info['table_index'] + 1}.")

            if all_formatted_sections:
                formatted_tables_text = "\n".join(all_formatted_sections)
                logger.info(f"Debug: Formatted table text length: {len(formatted_tables_text)}")
                if not formatted_tables_text.strip():
                    logger.warning("Debug: Formatted table text content is empty or only whitespace.")
            else:
                logger.info("No sections were generated from the extracted tables.")

        logger.debug(f"*******TEXTONLY {full_document_paragraph_text}");
        logger.debug(f"*******TABLEONLY {formatted_tables_text}");
        # --- Step 3: Combine Results ---
        if full_document_paragraph_text:
            # combined_document_content += "##################################################\n"
            # combined_document_content += "FULL DOCUMENT PARAGRAPH TEXT\n"
            # combined_document_content += "##################################################\n"
            combined_document_content += full_document_paragraph_text

        if formatted_tables_text:
            if combined_document_content:
                combined_document_content += "\n\n" # Add extra newlines for clear separation
            # combined_document_content += "==================================================\n"
            # combined_document_content += "EXTRACTED TABLE DATA (FORMATTED SECTIONS)\n"
            # combined_document_content += "==================================================\n"
            combined_document_content += formatted_tables_text

        logger.info(f"Debug: Combined content length before return: {len(combined_document_content)}")
        logger.debug(f"*******COBINED {combined_document_content}");
        if not combined_document_content.strip():
            logger.warning("Debug: Final combined content is empty or only whitespace before return.")
        # return full_document_paragraph_text
        return combined_document_content