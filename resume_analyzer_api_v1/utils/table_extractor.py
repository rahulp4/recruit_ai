# DOCX_FILE_PATH = "/Users/rahulpoddar/my-work/project_resume/resume_structurer_llm/sample_data/priyasharma.docx" # <--- IMPORTANT: Change this to your DOCX path!
import pandas as pd
from docx import Document
import io
import logging

# ---
## Configuration
# ---

# Define the path to your DOCX document for initial loading into a stream.
# Make sure to replace 'your_document.docx' with the actual path to your file.
# DOCX_FILE_PATH = "/Users/rahulpoddar/my-work/project_resume/resume_structurer_llm/sample_data/priyasharma.docx" # <--- IMPORTANT: Change this to your DOCX path!
DOCX_FILE_PATH = "/Users/rahulpoddar/my-work/project_resume/resume_structurer_llm/sample_data/Rahul_Poddar_V4.docx" # <--- IMPORTANT: Change this to your DOCX path!
# ---
## Logging Setup
# ---
# FIX: Removed the repeated 'format' keyword argument.
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ---
## Function to Extract Tables from DOCX
# ---

def extract_tables_from_docx(docx_stream):
    """
    Extracts tables from a DOCX document provided as a byte stream.

    Args:
        docx_stream (io.BytesIO or similar file-like object):
            A stream containing the byte content of the DOCX file.

    Returns:
        list: A list of dictionaries, where each dictionary contains:
              'table_index' (int) and 'dataframe' (pd.DataFrame) for an extracted table.
    """
    all_extracted_dfs = []
    try:
        # Ensure stream is at the start for this function
        docx_stream.seek(0)
        document = Document(docx_stream)
        logger.info("Successfully loaded DOCX document from stream for table extraction.")

        if not document.tables:
            logger.info("No tables found in the document.")
            return []

        for table_idx, table in enumerate(document.tables):
            logger.info(f"Processing Table {table_idx + 1}...")
            data = []
            keys = None

            # Iterate over rows
            for i, row in enumerate(table.rows):
                text_cells = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    # Assume the first row is the header
                    keys = text_cells
                else:
                    data.append(text_cells)

            if keys:
                # Ensure all data rows have the same number of columns as headers
                if data:
                    # Pad rows with None if they are shorter than the header, truncate if longer
                    max_cols = len(keys)
                    padded_data = [row + [None] * (max_cols - len(row)) if len(row) < max_cols else row[:max_cols] for row in data]
                    df = pd.DataFrame(padded_data, columns=keys)
                else: # Table has headers but no data rows
                    df = pd.DataFrame(columns=keys) # Create empty DataFrame with headers


                all_extracted_dfs.append({
                    "table_index": table_idx,
                    "dataframe": df
                })
                logger.info(f"  Extracted Table {table_idx + 1}.")
            else:
                logger.warning(f"  Table {table_idx + 1} has no header row or no content.")

    except Exception as e:
        logger.error(f"An unexpected error occurred during DOCX stream processing for tables: {e}", exc_info=True)

    return all_extracted_dfs

# ---
## Function to Extract Text from DOCX (only top-level paragraphs)
# ---

def extract_text_from_docx(docx_file_stream):
    """
    Extracts all textual content from a .docx file stream,
    focusing only on top-level paragraphs (not tables).
    """
    try:
        # Ensure stream is at the start for this function
        docx_file_stream.seek(0)
        document = Document(docx_file_stream)
        logger.info("Successfully loaded DOCX document from stream for paragraph text extraction.")
        full_text = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error extracting paragraph text from DOCX: {e}", exc_info=True)
        raise # Re-raise the exception if you want calling code to handle it

# ---
## Function to Format DataFrame Rows into Sections
# ---

def format_dataframe_rows_to_sections(df, date_column_name='Date Range'):
    """
    Formats each row of a DataFrame into a structured "block of sections,"
    starting with a date range.

    Args:
        df (pd.DataFrame): The DataFrame containing the tabular data.
        date_column_name (str): The name of the column containing the date range.

    Returns:
        list: A list of strings, where each string is a formatted section block.
    """
    formatted_sections = []
    if df.empty:
        logger.info("  DataFrame is empty, no sections to format.")
        return []

    # Attempt to find a suitable date column if the specified one doesn't exist
    if date_column_name not in df.columns:
        logger.warning(f"  Date column '{date_column_name}' not found. Looking for alternatives.")
        # Common alternative names
        possible_date_cols = ['Date', 'Period', 'Duration', 'Time Frame', 'Valid From', 'Invoice Date', 'Start Date']
        found_date_col = None
        for col in possible_date_cols:
            if col in df.columns:
                date_column_name = col
                found_date_col = True
                logger.info(f"  Using '{date_column_name}' as the date column.")
                break
        if not found_date_col:
            logger.warning("  Could not find a suitable date column. 'Date Range' will be 'N/A Date Range'.")
            date_column_name = None # Indicate that no suitable date column was found

    for index, row in df.iterrows():
        # Get the date range, default to 'N/A' if column not found or value is missing
        if date_column_name and date_column_name in row:
            date_range = str(row[date_column_name]).strip() if pd.notna(row[date_column_name]) else "N/A Date Range"
        else:
            date_range = "N/A Date Range"

        # Dynamically add all other columns to the section
        other_details = []
        for col_name, value in row.items():
            # Skip the date column itself, and any empty string column headers that might appear
            if col_name != date_column_name and col_name:
                # Basic cleaning for display
                clean_value = str(value).strip() if pd.notna(value) else "N/A"
                other_details.append(f"**{col_name}:** {clean_value}")

        details_block = "\n".join(other_details)

        section_block = f"""
---
## Period: {date_range}

{details_block}

"""
        formatted_sections.append(section_block)
    return formatted_sections

# ---
## Main Processing Function (extract_text_from_docxv2)
# ---

def extract_text_from_docxv2(docx_content_stream):
    """
    Extracts and combines both paragraph text and formatted table data from a DOCX stream.

    Args:
        docx_content_stream (io.BytesIO): A stream containing the byte content of the DOCX file.

    Returns:
        str: The combined string of paragraph text followed by formatted table sections.
    """
    logger.info("Starting extract_text_from_docxv2 for comprehensive content extraction.")

    full_document_paragraph_text = ""
    formatted_tables_text = ""
    combined_document_content = ""

    # --- Step 1: Extract General Paragraph Text ---
    try:
        # This function handles seeking to 0 itself
        full_document_paragraph_text = extract_text_from_docx(docx_content_stream)
        logger.info(f"Debug: Paragraph text length: {len(full_document_paragraph_text)}")
        if not full_document_paragraph_text.strip():
            logger.warning("Debug: Paragraph text content is empty or only whitespace.")
    except Exception as e:
        logger.error(f"Failed to extract full paragraph text from DOCX within extract_text_from_docxv2: {e}")

    # --- Step 2: Extract and Process Tables ---
    # This function also handles seeking to 0 itself
    extracted_tables_info = extract_tables_from_docx(docx_content_stream)

    if not extracted_tables_info:
        logger.info("No tables were successfully extracted by extract_tables_from_docx.")
    else:
        logger.info(f"Successfully extracted {len(extracted_tables_info)} tables by extract_tables_from_docx.")
        all_formatted_sections = []

        for table_info in extracted_tables_info:
            current_df = table_info['dataframe']
            table_idx = table_info['table_index']

            logger.info(f"Processing Table {table_idx + 1} for Section Formatting.")
            if current_df.empty:
                logger.info(f"Table {table_idx + 1} is empty after extraction. Skipping section formatting.")
                continue

            logger.debug(f"Original DataFrame Head for Table {table_idx + 1}:\n{current_df.head().to_string()}")

            formatted_table_sections = format_dataframe_rows_to_sections(
                current_df,
                date_column_name='Date Range' # Adjust this if your date column has a different name!
            )
            all_formatted_sections.extend(formatted_table_sections)
            logger.info(f"Generated {len(formatted_table_sections)} sections for Table {table_idx + 1}.")

        if all_formatted_sections:
            formatted_tables_text = "\n".join(all_formatted_sections)
            logger.info(f"Debug: Formatted table text length: {len(formatted_tables_text)}")
            if not formatted_tables_text.strip():
                logger.warning("Debug: Formatted table text content is empty or only whitespace.")
        else:
            logger.info("No sections were generated from the extracted tables.")

    # --- Step 3: Combine Results ---
    # Append paragraph text first
    if full_document_paragraph_text:
        combined_document_content += "##################################################\n"
        combined_document_content += "FULL DOCUMENT PARAGRAPH TEXT\n"
        combined_document_content += "##################################################\n"
        combined_document_content += full_document_paragraph_text

    # Append formatted tables next, with a separator if there was previous content
    if formatted_tables_text:
        if combined_document_content: # Add newlines if there's already paragraph text
            combined_document_content += "\n\n" # Add extra newlines for clear separation
        combined_document_content += "==================================================\n"
        combined_document_content += "EXTRACTED TABLE DATA (FORMATTED SECTIONS)\n"
        combined_document_content += "==================================================\n"
        combined_document_content += formatted_tables_text

    logger.info(f"Debug: Combined content length before return: {len(combined_document_content)}")
    if not combined_document_content.strip():
        logger.warning("Debug: Final combined content is empty or only whitespace before return.")

    return combined_document_content

# ---
## Main Execution Block
# ---

# if __name__ == "__main__":
#     logger.info("Script execution started from __main__.")

#     docx_stream_to_process = None
#     try:
#         with open(DOCX_FILE_PATH, 'rb') as f: # Open in binary read mode
#             docx_bytes = f.read()
#             docx_stream_to_process = io.BytesIO(docx_bytes)
#         logger.info(f"Loaded DOCX file '{DOCX_FILE_PATH}' into an in-memory stream for main processing.")
#     except FileNotFoundError:
#         logger.error(f"DOCX file not found at '{DOCX_FILE_PATH}'. Please check the path and try again.")
#         exit() # Exit if the file isn't found
#     except Exception as e:
#         logger.error(f"Error loading DOCX file into stream in main: {e}", exc_info=True)
#         exit()

#     # Call the new main processing function
#     final_combined_output = extract_text_from_docxv2(docx_stream_to_process)
#     logger.info(f"OUT PUT - {final_combined_output}");
#     # Print the final combined content returned by extract_text_from_docxv2
#     # if final_combined_output.strip():
#     #     print("\n\n" + "*"*60)
#     #     print("COMBINED DOCUMENT CONTENT (PARAGRAPHS THEN TABLES):")
#     #     print("*"*60 + "\n")
#     #     # Print a truncated version if very long
#     #     print(final_combined_output[:4000] + ("\n..." if len(final_combined_output) > 4000 else ""))

#     #     # Optionally, save the combined output to a file
#     #     # with open("combined_document_content.txt", "w", encoding="utf-8") as f:
#     #     #     f.write(final_combined_output)
#     #     # logger.info("\nCombined document content saved to 'combined_document_content.txt'")
#     # else:
#     #     logger.info("\nNo combined content generated from the document (it was empty).")

#     logger.info("\nScript finished.")